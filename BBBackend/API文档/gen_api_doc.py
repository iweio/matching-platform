"""生成双智能体婚恋相亲平台 API 接口文档 (Word)"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_code_block(doc, text, font_size=Pt(9)):
    """添加代码块"""
    for i, line in enumerate(text.strip().split('\n')):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        if i == 0:
            p.paragraph_format.space_before = Pt(6)
        if i == len(text.strip().split('\n')) - 1:
            p.paragraph_format.space_after = Pt(6)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = font_size
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        p.paragraph_format.left_indent = Cm(0.8)

def add_api_section(doc, method, path, desc, req_body=None, resp_body=None, note=None):
    """添加一个 API 接口说明"""
    # 接口标题
    h = doc.add_heading(f'{desc}', level=3)
    # 方法和路径
    p = doc.add_paragraph()
    run_method = p.add_run(f'{method} ')
    run_method.bold = True
    run_method.font.color.rgb = RGBColor(0x00, 0x6B, 0xD6) if method == 'GET' else RGBColor(0xD6, 0x3A, 0x00)
    run_method.font.size = Pt(11)
    run_path = p.add_run(f'/api{path}')
    run_path.font.name = 'Consolas'
    run_path.font.size = Pt(11)

    if req_body:
        doc.add_paragraph('请求参数：', style='List Bullet')
        add_code_block(doc, req_body)

    if resp_body:
        doc.add_paragraph('响应示例：', style='List Bullet')
        add_code_block(doc, resp_body)

    if note:
        p = doc.add_paragraph()
        run = p.add_run(f'说明：{note}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True

def add_simple_table(doc, headers, rows):
    """添加简单表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('双智能体婚恋相亲平台')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('API 接口文档')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('版本 1.0.0  |  2026-05-11')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
doc.add_heading('目  录', level=1)
doc.add_paragraph('（在 Word 中按 Ctrl+A → F9 刷新目录域）')
doc.add_paragraph('插入方式：引用 → 目录 → 自动目录')
doc.add_page_break()

# ============================================================
# 1. 概述
# ============================================================
doc.add_heading('1. 概述', level=1)

doc.add_heading('1.1 基础信息', level=2)
add_simple_table(doc, ['项目', '说明'], [
    ['Base URL', 'http://<host>:8080'],
    ['字符编码', 'UTF-8'],
    ['Content-Type', 'application/json'],
    ['认证方式', 'Bearer Token (JWT)，登录/注册外的接口需在 Header 中携带'],
    ['Token Header', 'Authorization: Bearer <token>'],
])

doc.add_heading('1.2 统一响应格式', level=2)
doc.add_paragraph('所有接口返回统一的 JSON 结构：')
add_code_block(doc, '''{
  "code": 200,        // 业务状态码，200 表示成功
  "msg": "success",   // 提示信息
  "data": { ... }     // 业务数据，失败时为 null
}''')

doc.add_heading('1.3 业务错误码', level=2)
add_simple_table(doc, ['code', '说明'], [
    ['200', '成功'],
    ['401', '未登录或 Token 过期'],
    ['403', '无权访问（如未解锁查看对方资料）'],
    ['404', '资源不存在（用户/匹配/智能体）'],
    ['409', '业务冲突（重复操作）'],
    ['400', '参数校验失败'],
    ['500', '服务器内部错误'],
])

doc.add_page_break()

# ============================================================
# 2. 认证模块 /api/auth
# ============================================================
doc.add_heading('2. 认证模块  /api/auth', level=1)

add_api_section(doc, 'POST', '/auth/register', '2.1 用户注册',
    req_body='''{
  "phone": "13800138000",      // 必填，11位手机号
  "nick": "小明",               // 必填，2-10个字符
  "password": "123456",        // 必填，6-32位
  "gender": 1,                 // 可选，1=男 2=女
  "age": 28                    // 可选
}''',
    resp_body='''{
  "code": 200,
  "msg": "注册成功",
  "data": {
    "userId": "u_abc123",
    "agentId": "agent_abc123",
    "token": "eyJhbGci...",
    "nick": "小明",
    "distillStatus": 0
  }
}''',
    note='注册成功后自动创建智能体并返回 JWT Token。distillStatus=0 表示尚未完成人格蒸馏。'
)

add_api_section(doc, 'POST', '/auth/login', '2.2 用户登录',
    req_body='''{
  "phone": "13800138000",      // 必填
  "password": "123456"         // 必填
}''',
    resp_body='''{
  "code": 200,
  "msg": "登录成功",
  "data": {
    "userId": "u_abc123",
    "agentId": "agent_abc123",
    "token": "eyJhbGci...",
    "nick": "小明",
    "distillStatus": 1
  }
}'''
)

add_api_section(doc, 'GET', '/auth/me', '2.3 获取当前用户信息',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": {
    "userId": "u_abc123",
    "agentId": "agent_abc123",
    "token": "eyJhbGci...",
    "nick": "小明",
    "distillStatus": 1
  }
}''',
    note='需携带 Authorization Header。用于验证 Token 是否有效或刷新用户状态。'
)

doc.add_page_break()

# ============================================================
# 3. 智能体模块 /api/agent
# ============================================================
doc.add_heading('3. 智能体模块  /api/agent', level=1)

add_api_section(doc, 'POST', '/agent/init', '3.1 创建智能体',
    req_body='''{
  "phone": "13800138000",      // 必填
  "nick": "小明",               // 必填，2-10个字符
  "gender": 1,                 // 必填，1=男 2=女
  "age": 28,                   // 必填
  "bottomLine": {}             // 可选，底线配置 JSON
}''',
    resp_body='''{
  "code": 200,
  "msg": "智能体创建成功",
  "data": {
    "userId": "u_abc123",
    "agentId": "agent_abc123"
  }
}''',
    note='注册时已自动创建智能体，此接口用于需要重新初始化或补建智能体的场景。'
)

doc.add_page_break()

# ============================================================
# 4. 用户模块 /api/user
# ============================================================
doc.add_heading('4. 用户模块  /api/user', level=1)

add_api_section(doc, 'POST', '/user/distill', '4.1 人格蒸馏',
    req_body='''{
  "speakStyle": "活泼",        // 必填，说话风格
  "character": "外向开朗",     // 必填，性格特征
  "loveStyle": "主动",         // 必填，恋爱模式
  "valuesView": {             // 必填，价值观（自由 JSON）
    "婚姻观": "认真",
    "家庭观": "传统"
  },
  "taboo": {                  // 必填，禁忌（自由 JSON）
    "接受不了": "抽烟"
  }
}''',
    resp_body='''{
  "code": 200,
  "msg": "人格蒸馏数据保存成功，模型生成中",
  "data": {
    "status": "completed",
    "modelVersion": "v_20260511_001"
  }
}''',
    note='需携带 Authorization Header。调用此接口后系统会异步调用 Agent Service 刷新模型。'
)

add_api_section(doc, 'GET', '/user/me?user_id=u_abc123', '4.2 获取个人信息',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": {
    "userId": "u_abc123",
    "nick": "小明",
    "gender": 1,
    "age": 28,
    "agentId": "agent_abc123",
    "distillStatus": 1,
    "matchCount": 5,
    "unlockCount": 2
  }
}''',
    note='user_id 可选，不传则返回当前登录用户。matchCount=累计匹配次数，unlockCount=解锁次数。'
)

add_api_section(doc, 'GET', '/user/other?user_id=u_abc123&match_id=m_xyz', '4.3 查看对方资料',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": {
    "userId": "u_def456",
    "nick": "小红",
    "gender": 2,
    "age": 26,
    "distillInfo": {
      "speak_style": "温柔",
      "character": "文静"
    },
    "matchDuration": "互动 3 天"
  }
}''',
    note='仅在双方均已同意解锁（unlock_flag=1 且 status=5）后方可查看。否则返回 403。'
)

add_api_section(doc, 'GET', '/user/conversations', '4.4 会话列表',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": [
    {
      "matchId": "m_xyz",
      "partnerUserId": "u_def456",
      "partnerNick": "小红",
      "partnerGender": 2,
      "lastMessage": "你好呀，很高兴认识你～",
      "lastTime": "21:30",
      "unlockFlag": 1
    }
  ]
}''',
    note='返回当前用户的所有匹配会话列表，按最近消息时间排序。unlockFlag=1 表示已解锁真人聊天。'
)

doc.add_page_break()

# ============================================================
# 5. 匹配模块 /api/match
# ============================================================
doc.add_heading('5. 匹配模块  /api/match', level=1)

doc.add_heading('5.1 匹配业务流程', level=2)
doc.add_paragraph(
    '匹配系统采用轮询队列机制：用户发起匹配后，系统查找同时段也在等待的异性用户。'
    '匹配成功后，双方智能体自动进行 30 轮对话，完成后生成匹配报告。'
    '双方均可查看报告并选择"同意解锁"或"拒绝"。双方都同意后，解锁真人聊天。'
)
add_simple_table(doc, ['状态码', '含义', '说明'], [
    ['0', '匹配中', '等待系统配对或智能体对话未开始'],
    ['1', '智能体对话中', '双方 AI 正在自动对话（30轮）'],
    ['2', '对话完成', '智能体对话结束，可查看报告'],
    ['3', '报告就绪', '匹配报告已生成'],
    ['4', '已拒绝', '任一方拒绝解锁'],
    ['5', '已解锁', '双方同意，开启真人聊天'],
])

add_api_section(doc, 'POST', '/match/start', '5.2 发起匹配',
    req_body='''{
  "userId": "u_abc123"        // 必填，当前用户ID
}''',
    resp_body='''// 匹配成功：
{
  "code": 200,
  "msg": "匹配已发起",
  "data": {
    "matchId": "m_abc123_def456",
    "status": 1,
    "queued": false
  }
}

// 进入等待队列：
{
  "code": 200,
  "msg": "匹配已发起",
  "data": {
    "queued": true,
    "message": "暂无可匹配用户，系统将为您自动寻找"
  }
}''',
    note='匹配成功后前端跳转到 /matching 页进行 SSE 实时监听智能体对话。若进入队列则每 3 秒重试。'
)

add_api_section(doc, 'GET', '/match/progress?match_id=m_xxx&user_id=u_xxx', '5.3 查询匹配进度',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": {
    "matchId": "m_abc123_def456",
    "status": 1,
    "partnerId": "u_def456",
    "chatRound": 15,
    "a_op": "none",
    "b_op": "none",
    "unlock_flag": 0
  }
}''',
    note='polling 方式查询。status 对照 5.1 节状态表。a_op/b_op 取值：none/agree/reject。'
)

add_api_section(doc, 'GET', '/match/waiting-status?user_id=u_xxx', '5.4 查询排队状态',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": {
    "queued": true,
    "message": "正在为您寻找合适的匹配对象..."
  }
}''',
    note='用于匹配排队期间的状态查询。'
)

add_api_section(doc, 'GET', '/match/chat-record?match_id=m_xxx&since_id=msg_0005', '5.5 查询智能体对话记录',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": {
    "records": [
      {
        "id": "msg_0006",
        "speaker": "agent_a",
        "content": "你说得很有道理～我也觉得真诚是最重要的",
        "timestamp": "2026-05-11 21:15:30"
      }
    ],
    "hasMore": false
  }
}''',
    note='since_id 可选，用于增量拉取。前端 SSE 断线后可用此接口补拉遗漏的消息。'
)

add_api_section(doc, 'GET', '/match/chat-stream?match_id=m_xxx&user_id=u_xxx', '5.6 SSE 实时对话流',
    note='''此接口返回 text/event-stream，非普通 JSON。前端使用 EventSource 连接。

事件类型：
  - message   → 新消息 { "id": "msg_0001", "speaker": "agent_a", "content": "你好", "timestamp": "..." }
  - done      → 对话结束
  - error     → 出错 { "msg": "错误描述" }

前端用法：
  const es = new EventSource(`/api/match/chat-stream?match_id=${matchId}&user_id=${uid}`);
  es.addEventListener("message", (e) => { const msg = JSON.parse(e.data); ... });
  es.addEventListener("done", () => { es.close(); ... });
  es.addEventListener("error", () => { ... });
''')

add_api_section(doc, 'POST', '/match/unlock', '5.7 解锁操作',
    req_body='''{
  "matchId": "m_abc123_def456",  // 必填
  "userId": "u_abc123",          // 必填
  "operation": "agree"           // 必填，agree=同意 reject=拒绝
}''',
    resp_body='''{
  "code": 200,
  "msg": "操作成功，等待对方确认",
  "data": {
    "unlockStatus": 0,
    "bothAgreed": false
  }
}

// 双方都同意后：
{
  "code": 200,
  "msg": "操作成功，等待对方确认",
  "data": {
    "unlockStatus": 5,
    "bothAgreed": true
  }
}''',
    note='双方都 agree 后 match.status 变为 5，unlock_flag=1，自动创建真人聊天会话。'
)

add_api_section(doc, 'GET', '/match/report?match_id=m_xxx&user_id=u_xxx', '5.8 查询匹配报告',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": {
    "matchId": "m_abc123_def456",
    "score": 78,
    "dimensions": {
      "emotion": 82,
      "value": 75,
      "communication": 80,
      "lifestyle": 72,
      "future": 81
    },
    "advantage": "双方沟通比较顺畅，核心价值观上较为一致...",
    "risk": "消费观念和生活方式上略有分歧...",
    "suggest": "整体匹配度不错，建议解锁真人聊天..."
  }
}''',
    note='score 为 0-100 综合评分。dimensions 包含五个维度的独立评分。'
)

add_api_section(doc, 'GET', '/match/history?user_id=u_xxx', '5.9 匹配历史记录',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": [
    {
      "match_id": "m_abc123_def456",
      "partner_id": "u_def456",
      "partner_nick": "小红",
      "status": 5,
      "a_op": "agree",
      "b_op": "agree",
      "unlock_flag": 1,
      "chat_round": 30,
      "create_time": "2026-05-10T14:30:00",
      "update_time": "2026-05-11T09:15:00",
      "score": 78,
      "advantage": "双方沟通比较顺畅..."
    }
  ]
}''',
    note='按创建时间倒序排列。'
)

doc.add_page_break()

# ============================================================
# 6. 真人聊天模块 /api/chat
# ============================================================
doc.add_heading('6. 真人聊天模块  /api/chat', level=1)

doc.add_paragraph('双方均同意解锁后（match.status=5），真人聊天功能开通。')

add_api_section(doc, 'POST', '/chat/send', '6.1 发送消息',
    req_body='''{
  "matchId": "m_abc123_def456",   // 必填
  "senderId": "u_abc123",         // 必填，消息发送者
  "content": "你好呀～终于可以聊天了！"  // 必填
}''',
    resp_body='''{
  "code": 200,
  "msg": "消息发送成功",
  "data": {
    "messageId": "msg_uuid_xxx",
    "timestamp": "2026-05-11 22:00:00"
  }
}'''
)

add_api_section(doc, 'GET', '/chat/list?match_id=m_xxx&user_id=u_xxx&page=1&page_size=20', '6.2 获取聊天记录',
    resp_body='''{
  "code": 200,
  "msg": "ok",
  "data": {
    "messages": [
      {
        "message_id": "msg_uuid_001",
        "sender_id": "u_abc123",
        "content": "你好呀～",
        "created_at": "2026-05-11 22:00:00"
      }
    ],
    "page": 1,
    "total": 45
  }
}''',
    note='支持分页，默认 page=1, page_size=20。按时间正序排列。'
)

doc.add_page_break()

# ============================================================
# 7. Agent 算法服务 /api/agent/algo (内部)
# ============================================================
doc.add_heading('7. Agent 算法服务  /api/agent/algo', level=1)
doc.add_paragraph(
    '以下接口由后端内部调用 Agent Service（Python FastAPI），前端不直接访问。'
    '此处列出供调试和对接参考。'
)

add_api_section(doc, 'POST', '/agent/algo/model-refresh', '7.1 模型刷新',
    req_body='''{
  "userId": "u_abc123",
  "agentId": "agent_abc123",
  "speakStyle": "活泼",
  "character": "外向开朗",
  "loveStyle": "主动",
  "valuesView": { "婚姻观": "认真" },
  "taboo": { "接受不了": "抽烟" }
}''',
    resp_body='''{
  "code": 200,
  "msg": "模型生成成功",
  "data": {
    "modelVersion": "v_20260511_001",
    "status": "completed"
  }
}'''
)

add_api_section(doc, 'POST', '/agent/algo/chat-start', '7.2 启动智能体对话',
    req_body='''{
  "matchId": "m_abc123_def456",
  "agentIdA": "agent_abc123",
  "agentIdB": "agent_def456",
  "roundLimit": 30
}''',
    resp_body='''{
  "code": 200,
  "msg": "对话已启动",
  "data": {
    "sessionId": "session_20260511_001",
    "chatStatus": "running"
  }
}'''
)

add_api_section(doc, 'POST', '/agent/algo/generate-report', '7.3 生成匹配报告',
    req_body='''{
  "sessionId": "session_20260511_001",
  "matchId": "m_abc123_def456"
}''',
    resp_body='''{
  "code": 200,
  "msg": "报告生成成功",
  "data": {
    "score": 78,
    "dimensions": {
      "emotion": 82, "value": 75,
      "communication": 80, "lifestyle": 72, "future": 81
    },
    "advantage": "双方沟通比较顺畅...",
    "risk": "消费观念和生活方式上略有分歧...",
    "suggest": "整体匹配度不错，建议解锁真人聊天..."
  }
}'''
)

add_api_section(doc, 'POST', '/agent/algo/risk-detect', '7.4 风险检测',
    req_body='''{
  "sessionId": "session_20260511_001",
  "chatRecord": [
    { "speaker": "agent_a", "content": "你好呀！" },
    { "speaker": "agent_b", "content": "很高兴认识你～" }
  ]
}''',
    resp_body='''{
  "code": 200,
  "msg": "检测完成",
  "data": {
    "riskTags": [],
    "riskScore": 5,
    "blockSuggest": "未检测到明显风险，对话正常进行"
  }
}''',
    note='riskScore 0-10 分，分数越高风险越大。blockSuggest 为空字符串表示无风险。'
)

doc.add_page_break()

# ============================================================
# 8. 完整业务流程
# ============================================================
doc.add_heading('8. 完整业务流程', level=1)

doc.add_heading('8.1 用户注册 → 人格蒸馏 → 匹配 → 解锁 → 真人聊天', level=2)

steps = [
    ('1. 注册/登录', 'POST /api/auth/register → 获得 userId + token + agentId'),
    ('2. 人格蒸馏', 'POST /api/user/distill → 设定说话风格/性格/恋爱模式/价值观/禁忌'),
    ('3. 发起匹配', 'POST /api/match/start → 系统寻找异性用户配对'),
    ('4. 智能体对话', 'SSE /api/match/chat-stream → 双方 AI 自动进行 30 轮对话'),
    ('5. 查看报告', 'GET /api/match/report → 查看五维匹配度评分'),
    ('6. 解锁操作', 'POST /api/match/unlock → agree/reject'),
    ('7. 真人聊天', 'POST /api/chat/send + GET /api/chat/list → 双方真人交流'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}  ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('8.2 前端路由映射', level=2)
add_simple_table(doc, ['路由', '页面', '说明'], [
    ['/login', 'LoginPage', '登录'],
    ['/register', 'RegisterPage', '注册'],
    ['/distill', 'DistillPage', '人格蒸馏表单'],
    ['/distill/result', 'DistillResultPage', '蒸馏结果/状态'],
    ['/match', 'MatchPage', '匹配主页 + 历史记录'],
    ['/matching', 'MatchingPage', 'SSE 实时智能体对话'],
    ['/report', 'ReportPage', '匹配报告五维图'],
    ['/unlock', 'UnlockPage', '解锁操作页'],
    ['/chat', 'ChatPage', '真人聊天'],
    ['/profile', 'ProfilePage', '个人中心'],
    ['/settings', 'SettingsPage', '设置'],
    ['*', 'NotFoundPage', '404'],
])

doc.add_page_break()

# ============================================================
# 9. 附录
# ============================================================
doc.add_heading('9. 附录', level=1)

doc.add_heading('9.1 HTTP 状态码速查', level=2)
add_simple_table(doc, ['状态码', '含义'], [
    ['200', '成功'],
    ['400', '请求参数校验失败'],
    ['401', '未认证'],
    ['403', '无权限'],
    ['404', '资源不存在'],
    ['409', '业务冲突'],
    ['500', '服务器内部错误'],
])

doc.add_heading('9.2 匹配状态流转', level=2)
doc.add_paragraph('''
0（匹配中）→ 1（智能体对话中）→ 2（对话完成）→ 3（报告就绪）
                                    ├── 4（任一方拒绝）
                                    └── 5（双方同意解锁 → 真人聊天）
''')

doc.add_heading('9.3 鉴权说明', level=2)
doc.add_paragraph(
    '除 POST /api/auth/register 和 POST /api/auth/login 外，所有接口均需在 HTTP Header 中携带：'
)
add_code_block(doc, 'Authorization: Bearer <jwt_token>')
doc.add_paragraph(
    'Token 在注册/登录响应中返回。后端通过 JWT 中间件解析 userId 注入到 request 上下文。'
    'Token 有效期默认 7 天。'
)

doc.add_heading('9.4 接口汇总', level=2)
add_simple_table(doc, ['方法', '路径', '说明', '认证'], [
    ['POST', '/api/auth/register', '用户注册', '无需'],
    ['POST', '/api/auth/login', '用户登录', '无需'],
    ['GET', '/api/auth/me', '当前用户信息', '需要'],
    ['POST', '/api/agent/init', '创建智能体', '无需'],
    ['POST', '/api/user/distill', '人格蒸馏', '需要'],
    ['GET', '/api/user/me', '个人信息', '需要'],
    ['GET', '/api/user/other', '对方资料', '需要'],
    ['GET', '/api/user/conversations', '会话列表', '需要'],
    ['POST', '/api/match/start', '发起匹配', '需要'],
    ['GET', '/api/match/progress', '匹配进度', '需要'],
    ['GET', '/api/match/waiting-status', '排队状态', '需要'],
    ['GET', '/api/match/chat-record', '智能体对话记录', '需要'],
    ['GET', '/api/match/chat-stream', 'SSE实时对话', '需要'],
    ['POST', '/api/match/unlock', '解锁操作', '需要'],
    ['GET', '/api/match/report', '匹配报告', '需要'],
    ['GET', '/api/match/history', '历史匹配', '需要'],
    ['POST', '/api/chat/send', '发送消息', '需要'],
    ['GET', '/api/chat/list', '聊天记录', '需要'],
    ['POST', '/api/agent/algo/model-refresh', '模型刷新', '内部'],
    ['POST', '/api/agent/algo/chat-start', '启动对话', '内部'],
    ['POST', '/api/agent/algo/generate-report', '生成报告', '内部'],
    ['POST', '/api/agent/algo/risk-detect', '风险检测', '内部'],
])

# ── 保存 ──
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, '双智能体婚恋相亲平台_API接口文档.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')
